"""
Главное окно OCR-приложения.
Предоставляет полный графический интерфейс со всеми элементами управления и функциональностью.
"""

import tempfile
from datetime import datetime
from pathlib import Path
from typing import Optional

from PySide6.QtWidgets import (
    QMainWindow,
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QSplitter,
    QToolBar,
    QMenuBar,
    QMenu,
    QStatusBar,
    QMessageBox,
    QFileDialog,
    QApplication,
    QLabel,
    QPushButton,
    QFrame,
)
from PySide6.QtGui import QAction, QIcon, QKeySequence
from PySide6.QtCore import Qt, QThread, Signal

from ..engine import OCREngine, OCRResult, TesseractEngine, KrakenEngine, TrOCREngine, CSLAVEngine
from ..preprocessing import PreprocessingPipeline, PipelineConfig
from .widgets import (
    ImagePreviewWidget,
    TextEditorWidget,
    EngineSelectorWidget,
    PreprocessingOptionsWidget,
    StatusBarWidget,
)


class OCRWorker(QThread):
    """Рабочий поток для обработки OCR."""
    
    finished = Signal(OCRResult)
    error = Signal(str)
    progress = Signal(int)  # 0-100
    processed_image_ready = Signal(str)  # Сигнал с путём к обработанному изображению
    
    def __init__(
        self,
        engine: OCREngine,
        image_path: Path,
        model_name: Optional[str] = None,
        pipeline: Optional[PreprocessingPipeline] = None,
    ):
        """Инициализировать рабочий поток OCR."""
        super().__init__()
        self.engine = engine
        self.image_path = image_path
        self.model_name = model_name
        self.pipeline = pipeline
        self.processed_image_path: Optional[Path] = None
    
    def run(self) -> None:
        """Запустить обработку OCR в фоновом потоке."""
        try:
            self.progress.emit(10)
            
            # Применить предобработку, если предоставлен конвейер
            if self.pipeline:
                self.progress.emit(20)
                # Создать временный файл для обработанного изображения
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                    tmp_path = Path(tmp.name)
                
                try:
                    processed_image = self.pipeline.process_file(self.image_path, tmp_path)
                    process_path = tmp_path
                    self.processed_image_path = tmp_path
                    
                    # Сигнализировать главному окну обновить предпросмотр выровненным изображением
                    self.processed_image_ready.emit(str(tmp_path))
                        
                except Exception:
                    process_path = self.image_path
            else:
                process_path = self.image_path
            
            self.progress.emit(50)
            
            # Запустить OCR
            result = self.engine.recognize(process_path, self.model_name)
            
            # НЕ удалять временный файл здесь - он может понадобиться для обновления UI
            # Файл будет удалён при создании нового временного файла или при закрытии приложения
            
            self.progress.emit(100)
            self.finished.emit(result)
            
        except Exception as e:
            self.error.emit(str(e))


class MainWindow(QMainWindow):
    """Главное окно приложения."""
    
    def __init__(self):
        """Инициализировать главное окно."""
        super().__init__()
        
        self.setWindowTitle("Historical Slavic OCR")
        self.setMinimumSize(1200, 800)
        
        # Инициализировать движки
        self._engines: dict[str, OCREngine] = {
            "tesseract": TesseractEngine(),
            "kraken": KrakenEngine(),
            "trocr": TrOCREngine(),
            "cslav": CSLAVEngine(),
        }
        
        self._current_engine: Optional[OCREngine] = None
        self._current_pipeline: Optional[PreprocessingPipeline] = None
        self._current_image_path: Optional[Path] = None
        self._last_export_dir: Optional[Path] = None
        self._processed_image_path: Optional[Path] = None  # Путь к обработанному изображению
        
        # Настроить UI
        self._setup_ui()
        self._setup_menu()
        self._setup_toolbar()
        self._connect_signals()
        
        # Инициализировать выбор движка
        self._update_engine_selector()
        
        # Установить статус
        self._status_bar.set_message("Готов - Загрузите изображение для начала работы")
    
    def _setup_ui(self) -> None:
        """Настроить пользовательский интерфейс."""
        # Центральный виджет
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QVBoxLayout(central_widget)
        main_layout.setContentsMargins(10, 10, 10, 10)
        main_layout.setSpacing(10)
        
        # Верхняя секция: элементы управления движком и предобработкой
        controls_layout = QHBoxLayout()
        controls_layout.setSpacing(20)
        
        # Выбор движка
        self._engine_selector = EngineSelectorWidget()
        self._engine_selector.setMaximumWidth(350)
        controls_layout.addWidget(self._engine_selector)
        
        # Параметры предобработки
        self._preprocessing_options = PreprocessingOptionsWidget()
        self._preprocessing_options.setMaximumWidth(300)
        controls_layout.addWidget(self._preprocessing_options)
        
        # Кнопка обработки
        button_layout = QVBoxLayout()
        
        self._process_btn = QPushButton("🔍 Распознать текст")
        self._process_btn.setMinimumHeight(40)
        self._process_btn.setEnabled(False)
        button_layout.addWidget(self._process_btn)
        
        self._clear_btn = QPushButton("Очистить")
        self._clear_btn.clicked.connect(self._clear_all)
        button_layout.addWidget(self._clear_btn)
        
        button_layout.addStretch()
        controls_layout.addLayout(button_layout)
        
        main_layout.addLayout(controls_layout)
        
        # Разделитель для изображения и текста
        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.setHandleWidth(5)
        
        # Слева: предпросмотр изображения
        self._image_preview = ImagePreviewWidget()
        splitter.addWidget(self._image_preview)
        
        # Справа: текстовый редактор
        self._text_editor = TextEditorWidget()
        splitter.addWidget(self._text_editor)
        
        # Установить начальные размеры (50/50)
        splitter.setSizes([600, 600])
        
        main_layout.addWidget(splitter, 1)
        
        # Строка состояния внизу
        self._status_bar = StatusBarWidget()
        main_layout.addWidget(self._status_bar)
    
    def _setup_menu(self) -> None:
        """Настроить меню."""
        menubar = self.menuBar()
        
        # Меню Файл
        file_menu = menubar.addMenu("&Файл")
        
        open_action = QAction("&Открыть изображение", self)
        open_action.setShortcut(QKeySequence.StandardKey.Open)
        open_action.triggered.connect(self._open_image)
        file_menu.addAction(open_action)
        
        file_menu.addSeparator()
        
        export_txt_action = QAction("Экспортировать как &TXT", self)
        export_txt_action.triggered.connect(lambda: self._export_text("txt"))
        file_menu.addAction(export_txt_action)
        
        export_docx_action = QAction("Экспортировать как &DOCX", self)
        export_docx_action.triggered.connect(lambda: self._export_text("docx"))
        file_menu.addAction(export_docx_action)
        
        file_menu.addSeparator()
        
        exit_action = QAction("В&ыход", self)
        exit_action.setShortcut(QKeySequence.StandardKey.Quit)
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
        
        # Меню Правка
        edit_menu = menubar.addMenu("&Правка")
        
        copy_action = QAction("&Копировать", self)
        copy_action.setShortcut(QKeySequence.StandardKey.Copy)
        copy_action.triggered.connect(self._text_editor.copy)
        edit_menu.addAction(copy_action)
        
        paste_action = QAction("&Вставить", self)
        paste_action.setShortcut(QKeySequence.StandardKey.Paste)
        paste_action.triggered.connect(self._text_editor.paste)
        edit_menu.addAction(paste_action)
        
        edit_menu.addSeparator()
        
        select_all_action = QAction("Выделить &всё", self)
        select_all_action.setShortcut(QKeySequence.StandardKey.SelectAll)
        select_all_action.triggered.connect(self._text_editor.selectAll)
        edit_menu.addAction(select_all_action)
        
        # Меню Вид
        view_menu = menubar.addMenu("&Вид")
        
        zoom_in_action = QAction("У&величить", self)
        zoom_in_action.setShortcut(QKeySequence.StandardKey.ZoomIn)
        zoom_in_action.triggered.connect(self._image_preview._zoom_in)
        view_menu.addAction(zoom_in_action)
        
        zoom_out_action = QAction("У&меньшить", self)
        zoom_out_action.setShortcut(QKeySequence.StandardKey.ZoomOut)
        zoom_out_action.triggered.connect(self._image_preview._zoom_out)
        view_menu.addAction(zoom_out_action)
        
        reset_zoom_action = QAction("&Сбросить масштаб", self)
        reset_zoom_action.setShortcut(QKeySequence("Ctrl+0"))
        reset_zoom_action.triggered.connect(lambda: setattr(self._image_preview, '_zoom_factor', 1.0) or self._image_preview._update_display())
        view_menu.addAction(reset_zoom_action)
        
        # Меню Справка
        help_menu = menubar.addMenu("&Справка")
        
        about_action = QAction("&О программе", self)
        about_action.triggered.connect(self._show_about)
        help_menu.addAction(about_action)
    
    def _setup_toolbar(self) -> None:
        """Настроить панель инструментов."""
        toolbar = QToolBar("Главная панель")
        toolbar.setMovable(False)
        self.addToolBar(toolbar)
        
        open_action = QAction("📁 Открыть", self)
        open_action.setToolTip("Открыть файл изображения")
        open_action.triggered.connect(self._open_image)
        toolbar.addAction(open_action)
        
        toolbar.addSeparator()
        
        process_action = QAction("🔍 Распознать", self)
        process_action.setToolTip("Распознать текст на изображении")
        process_action.triggered.connect(self._run_ocr)
        toolbar.addAction(process_action)
        
        toolbar.addSeparator()
        
        export_action = QAction("💾 Экспорт", self)
        export_action.setToolTip("Экспортировать распознанный текст")
        export_action.triggered.connect(lambda: self._export_text("txt"))
        toolbar.addAction(export_action)
    
    def _connect_signals(self) -> None:
        """Подключить сигналы виджетов."""
        # Выбор движка
        self._engine_selector.engine_changed.connect(self._on_engine_changed)
        self._engine_selector.model_changed.connect(self._on_model_changed)
        
        # Параметры предобработки
        self._preprocessing_options.options_changed.connect(self._on_preprocessing_changed)
        
        # Кнопки
        self._process_btn.clicked.connect(self._run_ocr)
        
        # Предпросмотр изображения
        self._image_preview.image_loaded.connect(self._on_image_loaded)
    
    def _update_engine_selector(self) -> None:
        """Обновить выбор движка доступными движками."""
        engines = [
            (name, engine.is_ready())
            for name, engine in self._engines.items()
        ]
        self._engine_selector.populate_engines(engines)
        
        # Выбрать первый доступный движок
        for name, available in engines:
            if available:
                self._on_engine_changed(name)
                break
    
    def _on_engine_changed(self, engine_name: str) -> None:
        """Обработать изменение выбора движка."""
        if engine_name not in self._engines:
            return
        
        engine = self._engines[engine_name]
        self._current_engine = engine
        
        if engine.is_ready():
            models = engine.get_available_models()
            self._engine_selector.populate_models(models)
            self._engine_selector.set_status("Готов", True)
        else:
            self._engine_selector.populate_models([])
            self._engine_selector.set_status("Недоступен", False)
    
    def _on_model_changed(self, model_name: str) -> None:
        """Обработать изменение выбора модели."""
        pass  # Модель используется во время выполнения OCR
    
    def _on_preprocessing_changed(self) -> None:
        """Обработать изменение параметров предобработки."""
        config_dict = self._preprocessing_options.get_config()
        
        # Сопоставить с PipelineConfig
        if all(not v for v in config_dict.values()):
            self._current_pipeline = None
        else:
            from ..preprocessing import ProcessorConfig
            
            config = PipelineConfig(
                denoise=ProcessorConfig(enabled=config_dict["denoise"]),
                clahe=ProcessorConfig(enabled=config_dict["clahe"]),
                deskew=ProcessorConfig(enabled=config_dict["deskew"]),
                contrast=ProcessorConfig(enabled=config_dict["contrast"]),
                binarization=ProcessorConfig(enabled=config_dict["binarization"]),
            )
            self._current_pipeline = PreprocessingPipeline(config)
    
    def _on_image_loaded(self, image_path: str) -> None:
        """Обработать событие загрузки изображения."""
        self._current_image_path = Path(image_path)
        self._process_btn.setEnabled(True)
        self._status_bar.set_message(f"Загружено: {Path(image_path).name}")
        self._status_bar.set_info(f"Размер: {self._image_preview._original_pixmap.width()}x{self._image_preview._original_pixmap.height()}px")
    
    def _open_image(self) -> None:
        """Открыть диалог выбора файла изображения."""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Открыть изображение",
            str(self._last_export_dir or Path.home()),
            "Файлы изображений (*.png *.jpg *.jpeg *.bmp *.tiff *.tif *.gif);;Все файлы (*)",
        )
        
        if file_path:
            self._current_image_path = Path(file_path)
            if self._image_preview.load_image(self._current_image_path):
                self._process_btn.setEnabled(True)
                self._status_bar.set_message(f"Загружено: {self._current_image_path.name}")
            else:
                QMessageBox.critical(self, "Ошибка", f"Не удалось загрузить изображение: {file_path}")
    
    def _run_ocr(self) -> None:
        """Запустить обработку OCR."""
        if not self._current_engine or not self._current_engine.is_ready():
            QMessageBox.warning(self, "Предупреждение", "Нет доступного OCR-движка")
            return
        
        # Использовать обработанное изображение если оно есть (для повторного распознавания), иначе оригинальное
        image_to_use = self._processed_image_path if self._processed_image_path and self._processed_image_path.exists() else self._current_image_path
        
        if not image_to_use:
            QMessageBox.warning(self, "Предупреждение", "Изображение не загружено")
            return
        
        # Отключить элементы управления во время обработки
        self._process_btn.setEnabled(False)
        self._process_btn.setText("Обработка...")
        self._status_bar.set_message("Обработка...")
        self._status_bar.set_progress(0, show=True)
        
        # Получить выбранную модель
        model_name = self._engine_selector.get_selected_model()
        
        # Создать рабочий поток - всегда используем оригинальное изображение для конвейера
        self._worker = OCRWorker(
            engine=self._current_engine,
            image_path=self._current_image_path,  # Всегда используем оригинал для предобработки
            model_name=model_name,
            pipeline=self._current_pipeline,
        )
        
        self._worker.finished.connect(self._on_ocr_finished)
        self._worker.error.connect(self._on_ocr_error)
        self._worker.progress.connect(self._status_bar.set_progress)
        self._worker.processed_image_ready.connect(self._on_processed_image_ready)
        
        self._worker.start()
    
    def _on_processed_image_ready(self, processed_image_path: str) -> None:
        """Обработать готовность обработанного изображения для обновления предпросмотра."""
        # Сохранить путь к обработанному изображению
        self._processed_image_path = Path(processed_image_path)
        # Обновить предпросмотр выровненным/обработанным изображением
        self._image_preview.load_image(self._processed_image_path)
    
    def _on_ocr_finished(self, result: OCRResult) -> None:
        """Обработать завершение OCR."""
        # Повторно включить элементы управления
        self._process_btn.setEnabled(True)
        self._process_btn.setText("🔍 Распознать текст")
        self._status_bar.set_progress(0, show=False)
        
        # Отобразить результат
        metadata = result.get_formatted_text().split('\n\n')[0] if '\n\n' in result.get_formatted_text() else ""
        self._text_editor.set_result_text(result.text, metadata)
        
        # Обновить статус
        if result.has_warnings():
            warning_text = "; ".join(result.warnings)
            self._status_bar.set_message(f"Завершено с предупреждениями: {warning_text}")
        else:
            self._status_bar.set_message("Распознавание успешно завершено")
        
        # Показать информацию
        info_parts = []
        if result.confidence > 0:
            info_parts.append(f"Достоверность: {result.confidence:.1%}")
        if result.processing_time_ms > 0:
            info_parts.append(f"Время: {result.processing_time_ms:.1f}мс")
        if info_parts:
            self._status_bar.set_info(" | ".join(info_parts))
    
    def _on_ocr_error(self, error_msg: str) -> None:
        """Обработать ошибку OCR."""
        self._process_btn.setEnabled(True)
        self._process_btn.setText("🔍 Распознать текст")
        self._status_bar.set_progress(0, show=False)
        
        QMessageBox.critical(self, "Ошибка OCR", f"Не удалось распознать текст:\n{error_msg}")
        self._status_bar.set_message("Ошибка при распознавании")
    
    def _export_text(self, format: str) -> None:
        """Экспортировать распознанный текст в файл."""
        text = self._text_editor.get_text()
        
        if not text.strip():
            QMessageBox.information(self, "Информация", "Нет текста для экспорта")
            return
        
        # Определить имя файла по умолчанию
        if self._current_image_path:
            default_name = self._current_image_path.stem
        else:
            default_name = "ocr_result"
        
        # Получить путь сохранения
        if format == "txt":
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Экспортировать как TXT",
                str(self._last_export_dir / f"{default_name}.txt") if self._last_export_dir else f"{default_name}.txt",
                "Текстовые файлы (*.txt);;Все файлы (*)",
            )
            if file_path:
                try:
                    with open(file_path, 'w', encoding='utf-8') as f:
                        f.write(text)
                    self._last_export_dir = Path(file_path).parent
                    self._status_bar.set_message(f"Экспортировано в {file_path}")
                except Exception as e:
                    QMessageBox.critical(self, "Ошибка", f"Не удалось экспортировать: {e}")
        
        elif format == "docx":
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Экспортировать как DOCX",
                str(self._last_export_dir / f"{default_name}.docx") if self._last_export_dir else f"{default_name}.docx",
                "Документы Word (*.docx);;Все файлы (*)",
            )
            if file_path:
                try:
                    from docx import Document
                    
                    doc = Document()
                    
                    # Добавить метаданные как абзац с меньшим шрифтом
                    lines = text.split('\n')
                    metadata_lines = []
                    content_start = 0
                    
                    for i, line in enumerate(lines):
                        if line.startswith(('Engine:', 'Model:', 'Confidence:', 'Processing time:', 'Движок:', 'Модель:', 'Достоверность:', 'Время обработки:')):
                            metadata_lines.append(line)
                            content_start = i + 1
                        else:
                            break
                    
                    if metadata_lines:
                        meta_para = doc.add_paragraph('\n'.join(metadata_lines))
                        for run in meta_para.runs:
                            run.font.size = doc.shared_styles['Normal'].font.size - 2
                    
                    # Добавить основное содержимое
                    content = '\n'.join(lines[content_start:])
                    if content.strip():
                        doc.add_paragraph(content)
                    
                    doc.save(file_path)
                    self._last_export_dir = Path(file_path).parent
                    self._status_bar.set_message(f"Экспортировано в {file_path}")
                    
                except ImportError:
                    QMessageBox.warning(
                        self, 
                        "Отсутствует зависимость", 
                        "python-docx не установлен. Установите: pip install python-docx"
                    )
                except Exception as e:
                    QMessageBox.critical(self, "Ошибка", f"Не удалось экспортировать: {e}")
    
    def _clear_all(self) -> None:
        """Очистить всё содержимое."""
        self._image_preview.clear()
        self._text_editor.clear()
        self._current_image_path = None
        self._processed_image_path = None  # Очистить путь к обработанному изображению
        self._process_btn.setEnabled(False)
        self._status_bar.reset()
        self._status_bar.set_message("Готов - Загрузите изображение для начала работы")
    
    def _show_about(self) -> None:
        """Показать диалог о программе."""
        QMessageBox.about(
            self,
            "О Historical Slavic OCR",
            "Приложение Historical Slavic OCR\n\n"
            "Инструмент для распознавания исторических славянских текстов с использованием нескольких OCR-движков:\n"
            "• Tesseract OCR — для печатных текстов\n"
            "• Kraken OCR — специализированный для исторических документов\n"
            "• TrOCR — продвинутое распознавание на основе трансформеров\n\n"
            "Возможности:\n"
            "• Поддержка нескольких движков с автоматическим переключением\n"
            "• Расширенная предобработка изображений\n"
            "• Поддержка различных славянских языков\n"
            "• Экспорт в форматы TXT и DOCX\n\n"
            "© 2024"
        )
